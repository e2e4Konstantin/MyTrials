from docx import Document
from docx.shared import RGBColor

document = Document()
document.add_paragraph().add_run("This run will be red").font.color.rgb = RGBColor(100,0,0)
document.add_paragraph("This paragraph will be the default color, so black probably")
document.save("style.docx")