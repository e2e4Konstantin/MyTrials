from docx import Document
import re
document = Document()
paragraph = document.add_paragraph('Lorem ipsum dolor sit amet.')
section = document.sections[0]
footer=section.footer
paragraph = footer.paragraphs[0]
paragraph.text = "Hopefully this is a footer"
out='./word/template_test.docx'
document.save(out)