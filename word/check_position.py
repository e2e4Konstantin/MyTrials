
import docx

from docx import Document
from docx.shared import Cm, Inches, Pt, RGBColor

# def get_element_position(doc: Document, paragraph: docx.text.paragraph.Paragraph) -> tuple:
    
#     page_height = doc.sections[-1].header._element.get('height')
#     print(f"{page_height=}")
#     element_top = paragraph.paragraph_format.paragraph_format_pr.top
#     element_left = paragraph.paragraph_format.paragraph_format_pr.left
#     print(element_top, element_left)
#     return (element_top, element_left)


doc = Document('./word/Limonad.docx')
section = doc.sections[0]
print("Default Orientation:", section.orientation)

for p in doc.element.xpath("./w:body/w:p[w:r/w:lastRenderedPageBreak]"):
    print(p.text)




p = doc.paragraphs[6]
print(type(p))
print(p.text)
br_counter = 0
for run in p.runs:
    if run._element.br_lst:             
        for br in run._element.br_lst:
            br_counter+=1
            print(br.type)

print(br_counter)
# print(run._element.xml)
print(run.contains_page_break)
print(run.iter_inner_content())
# 
print(p.contains_page_break)
print(p.rendered_page_breaks)
# print(p._p.xml)
formatting = p.paragraph_format
print('Не отрывать от следующего абзаца:', formatting.keep_with_next)
print('Не разрывать абзац:', formatting.keep_together)
print('Абзац с новой страницы:', formatting.page_break_before)
print('Запрет висячих строк:', formatting.widow_control)
print()

for run in p.runs:
    if run.italic:
        print(run.text)


for run in p.runs:
    if 'lastRenderedPageBreak' in run._element.xml:  
        print( 'soft page break found at run:', run.text[:20] )
    if 'w:br' in run._element.xml and 'type="page"' in run._element.xml:
        print ('hard page break found at run:', run.text[:20])



# for p in doc.paragraphs:
#      print(p.text)


# paragraph = doc.paragraphs[0]
# print(type(paragraph))
# print(paragraph.text)

# page_height = doc.sections[-1].header._element.get('height')
# print(f"{page_height=}")
# header_xml = doc.sections[-1].header.part._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}header')
# print(header_xml)


# position = get_element_position(doc, paragraph)
# print(f"Element position: ({position[0]}, {position[1]})")